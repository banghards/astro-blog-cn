from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
CONTENT_FILE = ROOT / "tools" / "safety_cases_content.txt"
TEMPLATE_FILE = Path(r"C:\Users\admin\Desktop\safety_cases_template.docx")
TEMP_OUTPUT_FILE = Path(r"C:\Users\admin\Desktop\13类安全生产高频亡人事故及案例（近年案例版v4）.docx")
SONGTI = "\u5b8b\u4f53"


def main() -> None:
    content = CONTENT_FILE.read_text(encoding="utf-8")
    doc = Document(str(TEMPLATE_FILE))

    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)

    first = True
    for line in content.splitlines():
        if not line.strip():
            doc.add_paragraph("")
            continue
        paragraph = doc.add_paragraph(line)
        paragraph.style = "Heading 1" if first else "Normal"
        first = False

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = SONGTI
            run._element.rPr.rFonts.set(qn("w:eastAsia"), SONGTI)
            run.font.size = Pt(16) if paragraph.style.name == "Heading 1" else Pt(12)

    doc.save(str(TEMP_OUTPUT_FILE))
    print(TEMP_OUTPUT_FILE)


if __name__ == "__main__":
    main()
